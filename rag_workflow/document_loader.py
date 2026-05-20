from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".html",
    ".htm",
    ".json",
    ".csv",
    ".xlsx",
    ".xls",
}


@dataclass(frozen=True)
class TextDocument:
    text: str
    metadata: dict


@dataclass(frozen=True)
class TextChunk:
    id: str
    text: str
    metadata: dict


def file_sha256(path: Path, block_size: int = 1024 * 1024) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(block_size), b""):
            digest.update(block)
    return digest.hexdigest()


def iter_supported_files(root: Path, extensions: set[str] | None = None) -> Iterable[Path]:
    allowed = {ext.lower() for ext in (extensions or SUPPORTED_EXTENSIONS)}
    if root.is_file():
        if root.suffix.lower() in allowed:
            yield root
        return

    for path in root.rglob("*"):
        if path.is_file() and path.suffix.lower() in allowed:
            yield path


def load_file(path: Path) -> TextDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    elif suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suffix in {".html", ".htm"}:
        text = _read_html(path)
    elif suffix == ".json":
        text = _read_json(path)
    elif suffix == ".csv":
        text = _read_csv(path)
    elif suffix in {".xlsx", ".xls"}:
        text = _read_excel(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    stat = path.stat()
    return TextDocument(
        text=clean_text(text),
        metadata={
            "source": str(path.resolve()),
            "file_name": path.name,
            "extension": suffix,
            "file_sha256": file_sha256(path),
            "modified_time": int(stat.st_mtime),
            "size_bytes": stat.st_size,
        },
    )


def chunk_document(
    document: TextDocument,
    *,
    chunk_size: int,
    chunk_overlap: int,
    min_chunk_length: int,
) -> list[TextChunk]:
    chunks = chunk_text(
        document.text,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        min_chunk_length=min_chunk_length,
    )
    source_hash = document.metadata["file_sha256"][:16]
    result: list[TextChunk] = []
    for index, text in enumerate(chunks):
        chunk_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]
        metadata = {
            **document.metadata,
            "chunk_index": index,
            "chunk_count": len(chunks),
        }
        result.append(
            TextChunk(
                id=f"{source_hash}:{index}:{chunk_hash}",
                text=text,
                metadata=metadata,
            )
        )
    return result


def chunk_text(
    text: str,
    *,
    chunk_size: int,
    chunk_overlap: int,
    min_chunk_length: int,
) -> list[str]:
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size")

    paragraphs = [part.strip() for part in text.split("\n\n") if part.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            chunks.extend(_sliding_chunks(paragraph, chunk_size, chunk_overlap))
            continue

        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            chunks.append(current.strip())
            current = _overlap_tail(current, chunk_overlap)
            current = f"{current}\n\n{paragraph}".strip() if current else paragraph

    if current:
        chunks.append(current.strip())

    return [chunk for chunk in chunks if len(chunk) >= min_chunk_length]


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if not blank:
                cleaned.append("")
            blank = True
            continue
        cleaned.append(line)
        blank = False
    return "\n".join(cleaned).strip()


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {page_number}]\n{page_text}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))
    return "\n\n".join(paragraphs + table_rows)


def _read_html(path: Path) -> str:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n")


def _read_json(path: Path) -> str:
    payload = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _read_csv(path: Path) -> str:
    frame = pd.read_csv(path)
    return frame.to_markdown(index=False)


def _read_excel(path: Path) -> str:
    sheets = pd.read_excel(path, sheet_name=None)
    parts = []
    for sheet_name, frame in sheets.items():
        parts.append(f"# Sheet: {sheet_name}\n\n{frame.to_markdown(index=False)}")
    return "\n\n".join(parts)


def _sliding_chunks(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(0, end - chunk_overlap)
    return [chunk for chunk in chunks if chunk]


def _overlap_tail(text: str, chunk_overlap: int) -> str:
    if chunk_overlap <= 0:
        return ""
    return text[-chunk_overlap:].strip()
